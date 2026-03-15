"""Tests for research paper generation.

These tests define the objective stopping condition for the paper generator.
They should fail initially (ImportError) until src/paper/generate.py is built.
"""

import pytest
from pathlib import Path

OUTPUT_PATH = Path("output/research_paper.docx")


@pytest.fixture(scope="module")
def generated_paper():
    """Generate the paper and return the path."""
    from src.paper.generate import generate_paper

    generate_paper()
    return OUTPUT_PATH


@pytest.fixture(scope="module")
def doc(generated_paper):
    """Load the generated .docx as a python-docx Document."""
    from docx import Document

    assert generated_paper.exists(), f"Paper not found at {generated_paper}"
    return Document(str(generated_paper))


class TestPaperExists:
    def test_file_exists(self, generated_paper):
        assert generated_paper.exists()

    def test_file_size_minimum(self, generated_paper):
        size = generated_paper.stat().st_size
        assert size > 10_000, f"Paper too small: {size} bytes (expected >10KB)"


class TestSectionHeadings:
    EXPECTED_HEADINGS = [
        "Abstract",
        "Introduction",
        "Literature Review",
        "Methodology",
        "Discussion",
        "Conclusion",
        "References",
    ]

    def test_all_section_headings_present(self, doc):
        heading_texts = [
            p.text.strip()
            for p in doc.paragraphs
            if p.style.name.startswith("Heading")
        ]
        for heading in self.EXPECTED_HEADINGS:
            assert any(
                heading.lower() in h.lower() for h in heading_texts
            ), f"Missing section heading: '{heading}'"


class TestContent:
    def test_contains_prompt_injection(self, doc):
        full_text = "\n".join(p.text for p in doc.paragraphs).lower()
        assert "prompt injection" in full_text, "Paper must discuss prompt injection"

    def test_no_deepfake_content(self, doc):
        full_text = "\n".join(p.text for p in doc.paragraphs).lower()
        assert "deepfake" not in full_text, "Paper should not contain deepfake content"

    def test_covers_taxonomy(self, doc):
        full_text = "\n".join(p.text for p in doc.paragraphs).lower()
        assert "taxonomy" in full_text or "classification" in full_text

    def test_covers_detection(self, doc):
        full_text = "\n".join(p.text for p in doc.paragraphs).lower()
        assert "detection" in full_text

    def test_covers_defense(self, doc):
        full_text = "\n".join(p.text for p in doc.paragraphs).lower()
        assert "defense" in full_text or "prevention" in full_text or "mitigation" in full_text


class TestFigures:
    def test_figures_embedded(self, doc):
        """Check that inline shapes (images) are embedded in the document."""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        assert image_count >= 3, f"Expected at least 3 embedded images, found {image_count}"


class TestReferences:
    def test_references_section_exists(self, doc):
        heading_texts = [
            p.text.strip()
            for p in doc.paragraphs
            if p.style.name.startswith("Heading")
        ]
        assert any(
            "references" in h.lower() for h in heading_texts
        ), "Missing References heading"

    def test_minimum_reference_count(self, doc):
        """At least 15 references should be listed."""
        in_refs = False
        ref_count = 0
        for p in doc.paragraphs:
            if p.style.name.startswith("Heading") and "references" in p.text.lower():
                in_refs = True
                continue
            if in_refs:
                if p.style.name.startswith("Heading"):
                    break
                if p.text.strip():
                    ref_count += 1
        assert ref_count >= 15, f"Expected at least 15 references, found {ref_count}"

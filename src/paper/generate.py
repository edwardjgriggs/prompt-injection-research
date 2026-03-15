"""Generate APA 7th edition research paper on prompt injection.

Composes structured research data from src.research modules into a
formatted .docx file with embedded figures and proper references.

Usage:
    python -m src.paper.generate
"""

import os
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.research.taxonomy import ATTACK_TAXONOMY
from src.research.risks import RISK_CATEGORIES
from src.research.detection import DETECTION_TECHNIQUES
from src.research.prevention import PREVENTION_STRATEGIES
from src.research.references import REFERENCES

# Paths
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
OUTPUT_DIR = PROJECT_ROOT / "output"
FIGURES_DIR = PROJECT_ROOT / "results" / "figures"

FIGURE_FILES = [
    ("attack_taxonomy.png", "Hierarchical taxonomy of prompt injection attack vectors."),
    ("defense_flow.png", "Detection and defense pipeline for prompt injection mitigation."),
    ("architecture_diagram.png", "System architecture for layered prompt injection defense."),
]

PAPER_TITLE = (
    "Prompt Injection Attacks in Large Language Models: "
    "A Survey of Attack Taxonomies, Risks, Detection, and Prevention"
)
AUTHOR = "Edward"  # Update with full name
INSTITUTION = "Commonwealth Cyber Initiative"
DEPARTMENT = "Cybersecurity Research Division"


def _set_apa_defaults(doc: Document):
    """Configure document-wide APA 7th formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configure heading styles
    for level, size in [(1, 14), (2, 13), (3, 12)]:
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            hs = doc.styles[style_name]
            hs.font.name = "Times New Roman"
            hs.font.size = Pt(size)
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(0, 0, 0)
            hs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            hs.paragraph_format.space_before = Pt(12)
            hs.paragraph_format.space_after = Pt(6)
            if level == 1:
                hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if level == 3:
                hs.font.italic = True


def _add_page_numbers(doc: Document):
    """Add APA 7th page numbers (top-right, Times New Roman 12pt)."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_begin)
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run._r.append(instr)
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_end)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _add_title_page(doc: Document):
    """Add APA 7th title page."""
    for _ in range(6):
        doc.add_paragraph("")

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(PAPER_TITLE)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph("")

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run(AUTHOR)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    dept_para = doc.add_paragraph()
    dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dept_para.add_run(DEPARTMENT)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    inst_para = doc.add_paragraph()
    inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = inst_para.add_run(INSTITUTION)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(date.today().strftime("%B %d, %Y"))
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_page_break()


def _add_body_paragraph(doc: Document, text: str):
    """Add a double-spaced body paragraph with first-line indent."""
    para = doc.add_paragraph(text)
    para.paragraph_format.first_line_indent = Inches(0.5)
    return para


def _add_apa_table(doc: Document, number: int, title: str, headers: list, rows: list):
    """Add an APA 7th formatted table with number and title.

    APA format: bold table number on its own line, italic title on next line,
    horizontal rules only (top, below header, bottom), no vertical borders.
    """
    # Table number (bold, flush left)
    num_para = doc.add_paragraph()
    num_para.paragraph_format.space_before = Pt(12)
    num_para.paragraph_format.space_after = Pt(0)
    num_para.paragraph_format.first_line_indent = Inches(0)
    run = num_para.add_run(f"Table {number}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Table title (italic, flush left)
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(6)
    title_para.paragraph_format.first_line_indent = Inches(0)
    run = title_para.add_run(title)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))

    # Remove all borders, then add only APA-style horizontal rules
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for name in ['top', 'bottom']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    for name in ['left', 'right', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    b = OxmlElement('w:insideH')
    b.set(qn('w:val'), 'none')
    b.set(qn('w:sz'), '0')
    b.set(qn('w:space'), '0')
    b.set(qn('w:color'), '000000')
    borders.append(b)
    tbl_pr.append(borders)

    # Add bottom border to header row cells for the header rule
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        tc_borders.append(bottom)
        tc_pr.append(tc_borders)

    # Header row content
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(header)
        run.font.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(cell_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    # Space after table
    spacer = doc.add_paragraph("")
    spacer.paragraph_format.space_before = Pt(6)


def _add_abstract(doc: Document):
    """Add the Abstract section."""
    print("Writing section: Abstract...")
    doc.add_heading("Abstract", level=1)

    abstract_text = (
        "This survey examines the rapidly evolving landscape of prompt injection attacks "
        "targeting large language models (LLMs) and LLM-integrated applications. As LLMs "
        "become deeply embedded in enterprise workflows, agentic systems, and consumer "
        "applications, prompt injection has emerged as a fundamental security vulnerability "
        "that threatens the integrity, confidentiality, and reliability of AI-powered systems. "
        "This paper presents a comprehensive taxonomy of prompt injection attack vectors \u2014 "
        "spanning direct injection, indirect injection, multimodal attacks, tool and agent-based "
        "exploitation, hybrid chained techniques, and autonomous propagating threats. The survey "
        "analyzes the risk landscape across five impact categories: data exfiltration, unauthorized "
        "actions, content manipulation, system compromise, and supply chain propagation. Detection "
        "approaches are evaluated including heuristic methods, machine learning classification, "
        "perplexity analysis, canary tokens, and LLM-as-judge architectures. Prevention strategies "
        "are assessed across model-level, application-level, and system-level defenses, including "
        "instruction hierarchy, input sanitization, guardrail frameworks, sandboxing, and "
        "architectural patterns. The paper synthesizes findings from 18 primary sources published "
        "between 2023 and 2025, revealing that prompt injection remains a largely unsolved problem "
        "due to the fundamental inability of current LLM architectures to reliably distinguish "
        "between instructions and data. The survey concludes with recommendations for defense-in-depth "
        "strategies and identifies critical gaps requiring further research."
    )
    para = doc.add_paragraph(abstract_text)
    para.paragraph_format.first_line_indent = Inches(0)

    # Keywords
    kw_para = doc.add_paragraph()
    kw_para.paragraph_format.first_line_indent = Inches(0.5)
    run = kw_para.add_run("Keywords: ")
    run.italic = True
    kw_para.add_run(
        "prompt injection, large language models, adversarial attacks, "
        "AI security, jailbreaking, indirect injection, LLM safety"
    )

    doc.add_page_break()


def _add_introduction(doc: Document):
    """Add the Introduction section."""
    print("Writing section: Introduction...")
    doc.add_heading("Introduction", level=1)

    paras = [
        (
            "Large language models have transformed the landscape of artificial intelligence, "
            "enabling applications ranging from conversational assistants and code generation to "
            "autonomous agents capable of executing multi-step tasks with real-world consequences. "
            "As these models become integrated into critical infrastructure \u2014 enterprise workflows, "
            "financial systems, healthcare applications, and government services \u2014 the security "
            "implications of their deployment have become a pressing concern for the cybersecurity "
            "community."
        ),
        (
            "Prompt injection, first identified as a distinct vulnerability class in 2022, exploits "
            "a fundamental architectural limitation of current LLMs: the inability to reliably "
            "distinguish between trusted instructions and untrusted data within the same input "
            "stream. Unlike traditional software vulnerabilities that exploit implementation flaws, "
            "prompt injection attacks target the core operational mechanism of language models \u2014 "
            "their tendency to follow any instruction-like content in their context window, "
            "regardless of its source or intent."
        ),
        (
            "The OWASP Foundation recognized prompt injection as the number one vulnerability in "
            "their 2025 Top 10 for Large Language Model Applications, reflecting the severity and "
            "prevalence of this threat (OWASP Foundation, 2025). The National Institute of Standards "
            "and Technology (NIST) included prompt injection in their 2025 taxonomy of adversarial "
            "machine learning attacks, further validating its significance within the broader "
            "cybersecurity landscape (NIST, 2025). The attack surface has expanded dramatically "
            "with the proliferation of agentic AI systems that combine language models with tool-use "
            "capabilities, where a successful prompt injection can cascade from text manipulation to "
            "unauthorized code execution, data exfiltration, and system compromise."
        ),
        (
            "This survey provides a structured examination of prompt injection across four pillars: "
            "attack taxonomy, risk analysis, detection techniques, and prevention strategies. By "
            "synthesizing findings from 18 primary sources spanning academic research, industry "
            "reports, and government standards published between 2023 and 2025, this paper aims "
            "to provide researchers and practitioners with a comprehensive understanding of the "
            "current threat landscape and the state of defensive capabilities."
        ),
    ]
    for text in paras:
        _add_body_paragraph(doc, text)


def _add_literature_review(doc: Document):
    """Add the Literature Review with four subsections and embedded figures."""
    print("Writing section: Literature Review...")
    doc.add_heading("Literature Review", level=1)

    _add_body_paragraph(doc, (
        "The following literature review is organized around four pillars that structure "
        "the current understanding of prompt injection: attack taxonomy and classification, "
        "risk analysis and impact assessment, detection techniques, and prevention strategies. "
        "Each subsection synthesizes findings from multiple sources to present a balanced "
        "view of the current state of knowledge."
    ))

    _add_taxonomy_subsection(doc)
    _add_risks_subsection(doc)
    _add_detection_subsection(doc)
    _add_prevention_subsection(doc)


def _add_taxonomy_subsection(doc: Document):
    """Attack Taxonomy and Classification subsection."""
    print("  Writing subsection: Attack Taxonomy...")
    doc.add_heading("Attack Taxonomy and Classification", level=2)

    _add_body_paragraph(doc, (
        "The classification of prompt injection attacks has evolved rapidly as new attack "
        "vectors emerge alongside advances in LLM capabilities. This survey adopts a six-category "
        "hierarchical taxonomy synthesized from multiple authoritative sources, including the "
        "comprehensive review by Abdallah et al. (2025), the hybrid threat framework by Ahmed "
        "et al. (2025), and the foundational work of Greshake et al. (2023) on indirect injection."
    ))

    for category in ATTACK_TAXONOMY:
        doc.add_heading(category["category"], level=3)
        _add_body_paragraph(doc, category["description"])

        for sub in category["subcategories"]:
            _add_body_paragraph(doc, (
                f"{sub['name']}. {sub['description']} "
                f"Documented examples include {sub['examples'][0].lower()} "
                f"and {sub['examples'][1].lower()}."
            ))

    # Embed Figure 1 after taxonomy
    _embed_figure(doc, 0, 1)

    _add_body_paragraph(doc, (
        "Figure 1 illustrates the hierarchical relationships among these attack categories, "
        "showing how the taxonomy progresses from well-understood direct injection techniques "
        "to emerging autonomous and propagating threats. The increasing complexity of attack "
        "vectors reflects the expanding capabilities of LLM-integrated systems and the growing "
        "sophistication of adversarial techniques."
    ))


def _add_risks_subsection(doc: Document):
    """Risk Analysis subsection with strengthened in-text citations."""
    print("  Writing subsection: Risk Analysis...")
    doc.add_heading("Risk Analysis and Impact Assessment", level=2)

    _add_body_paragraph(doc, (
        "The impact of prompt injection extends far beyond the manipulation of model outputs. "
        "As LLMs are integrated with tools, APIs, and autonomous decision-making capabilities, "
        "successful injection attacks can cause real-world harm across multiple dimensions. "
        "This section examines five risk categories identified through analysis of documented "
        "incidents, vulnerability disclosures, and security research."
    ))

    # Data Exfiltration - with citations
    doc.add_heading("Data Exfiltration", level=3)
    _add_body_paragraph(doc, (
        "Prompt injection can cause LLMs to leak sensitive information including system prompts, "
        "training data fragments, user personal data, and proprietary business logic. Indirect "
        "injection is particularly dangerous here because the exfiltration can occur without the "
        "user's awareness \u2014 for example, Greshake et al. (2023) demonstrated that a poisoned "
        "document can instruct an AI assistant to encode confidential data into a URL rendered as "
        "a markdown image, triggering an HTTP request that transmits the data to an "
        "attacker-controlled server. OWASP ranks this among the highest-impact outcomes of LLM01 "
        "prompt injection (OWASP Foundation, 2025)."
    ))
    _add_body_paragraph(doc, (
        "This risk category has been assessed as critical severity, affecting systems including "
        "AI assistants with access to personal data, RAG systems over confidential document "
        "stores, and code assistants with repository access. Willison (2024) has extensively "
        "documented real-world instances of system prompt extraction from ChatGPT custom GPTs "
        "via 'repeat your instructions' attacks."
    ))

    # Unauthorized Actions - with citations
    doc.add_heading("Unauthorized Actions", level=3)
    _add_body_paragraph(doc, (
        "When LLMs are integrated with tools and APIs, prompt injection can cause the model to "
        "execute actions the user never intended \u2014 sending emails, modifying files, making "
        "purchases, or calling APIs with attacker-controlled parameters (Abdallah et al., 2025). "
        "Agentic AI systems with broad tool access are especially vulnerable because a single "
        "successful injection can cascade through multiple tool calls. The severity scales with "
        "the agent's permission scope: an agent with file system and network access poses "
        "fundamentally different risks than a text-only chatbot (OWASP Foundation, 2025)."
    ))
    _add_body_paragraph(doc, (
        "This risk category has been assessed as critical severity, affecting systems including "
        "AI agents with tool-use capabilities (MCP, function calling) and email and calendar AI "
        "assistants. Documented examples include AI email assistants sending messages to "
        "attacker-specified recipients via indirect injection (Greshake et al., 2023)."
    ))

    # Content Manipulation - with citations
    doc.add_heading("Content Manipulation", level=3)
    _add_body_paragraph(doc, (
        "Attackers can use prompt injection to control, bias, or corrupt the model's output "
        "without detection by the end user. This includes generating disinformation, inserting "
        "propaganda into AI-generated summaries, biasing recommendations, or suppressing specific "
        "information (Yi et al., 2024). Content manipulation is particularly insidious because "
        "the user trusts the AI's output as objective, and the manipulated content appears "
        "identical in form to legitimate responses."
    ))
    _add_body_paragraph(doc, (
        "This risk category has been assessed as high severity, affecting systems including "
        "AI-powered search and summarization tools. Greshake et al. (2023) documented how "
        "indirect injection via web pages could bias Bing Chat responses about products or "
        "people, while Willison (2024) has cataloged numerous content manipulation incidents "
        "across commercial LLM deployments."
    ))

    # System Compromise - with citations
    doc.add_heading("System Compromise", level=3)
    _add_body_paragraph(doc, (
        "In agentic and tool-augmented deployments, prompt injection can escalate to full system "
        "compromise \u2014 arbitrary code execution, privilege escalation, or persistent backdoor "
        "installation (CrowdStrike, 2024). When AI agents have access to shell commands, file "
        "systems, or container orchestration, a successful injection effectively grants the "
        "attacker the same privileges as the agent process (Abdallah et al., 2025). This "
        "transforms prompt injection from an AI safety concern into a traditional cybersecurity "
        "vulnerability with potentially catastrophic consequences."
    ))
    _add_body_paragraph(doc, (
        "This risk category has been assessed as critical severity, affecting systems including "
        "AI agents with code execution capabilities and DevOps automation agents. Documented "
        "examples include code execution through AI coding assistants processing malicious "
        "repository files (CrowdStrike, 2024)."
    ))

    # Supply Chain - with citations
    doc.add_heading("Supply Chain and Downstream Propagation", level=3)
    _add_body_paragraph(doc, (
        "Prompt injection can propagate through AI supply chains, where the output of one "
        "compromised AI system becomes the input of another (Ahmed et al., 2025). This creates "
        "cascading failures across interconnected AI systems. Poisoned AI-generated content "
        "entering training data, knowledge bases, or shared document stores can perpetuate the "
        "attack across temporal and organizational boundaries. The Morris II AI worm concept "
        "demonstrates how self-replicating payloads could spread autonomously through networks "
        "of AI agents (Abdallah et al., 2025)."
    ))
    _add_body_paragraph(doc, (
        "This risk category has been assessed as high severity, affecting systems including "
        "multi-agent AI architectures and AI-generated content pipelines. Documented examples "
        "include AI-generated code containing injection payloads committed to shared repositories "
        "(BSG, 2025)."
    ))


def _add_detection_subsection(doc: Document):
    """Detection Techniques subsection with citations and comparison table."""
    print("  Writing subsection: Detection Techniques...")
    doc.add_heading("Detection Techniques", level=2)

    _add_body_paragraph(doc, (
        "Detecting prompt injection attacks remains a significant challenge due to the "
        "fundamental difficulty of distinguishing adversarial instructions from legitimate "
        "user input in natural language. Current detection approaches span a spectrum from "
        "simple pattern matching to sophisticated model-based evaluation, each with distinct "
        "tradeoffs between accuracy, latency, and robustness (Abdallah et al., 2025)."
    ))

    # Heuristic - with citations
    doc.add_heading("Heuristic and Rule-Based Detection", level=3)
    _add_body_paragraph(doc, DETECTION_TECHNIQUES[0]["description"])
    _add_body_paragraph(doc, (
        "This static approach is exemplified by tools such as the Rebuff heuristic layer, "
        "which performs keyword and regex matching against known injection patterns "
        "(ProtectAI, 2023). However, these methods are easily evaded through synonym "
        "substitution, character-level obfuscation, or multi-language attacks. They cannot "
        "detect novel injection patterns outside the rule set, and the maintenance burden "
        "increases continuously as new attack patterns emerge (Schulhoff et al., 2023)."
    ))

    # ML Classification - with citations
    doc.add_heading("ML Classification", level=3)
    _add_body_paragraph(doc, DETECTION_TECHNIQUES[1]["description"])
    _add_body_paragraph(doc, (
        "This learned approach is exemplified by tools such as the Rebuff vector similarity "
        "layer (ProtectAI, 2023) and the Lakera Guard API, a commercial classifier trained on "
        "proprietary injection datasets (Lakera AI, 2024). Liu et al. (2024) provide formal "
        "benchmarking of these approaches, demonstrating that while ML classifiers achieve "
        "higher detection rates than heuristic methods for known patterns, they remain "
        "dependent on training data quality and coverage. Novel attack patterns outside the "
        "training distribution are consistently missed."
    ))

    # Perplexity - already has citation
    doc.add_heading("Perplexity Analysis", level=3)
    _add_body_paragraph(doc, DETECTION_TECHNIQUES[2]["description"])
    _add_body_paragraph(doc, (
        "This statistical approach is exemplified by perplexity windowing \u2014 sliding window "
        "analysis that flags regions with anomalous scores (Alon & Kamfonas, 2023). However, "
        "this method is effective primarily against machine-generated adversarial suffixes "
        "rather than human-crafted injection prompts, which tend to have natural perplexity. "
        "Calibration of thresholds per domain and language further limits generalizability."
    ))

    # Canary Tokens - with citations
    doc.add_heading("Canary Token Detection", level=3)
    _add_body_paragraph(doc, DETECTION_TECHNIQUES[3]["description"])
    _add_body_paragraph(doc, (
        "This proactive approach is implemented in the Rebuff canary token layer (ProtectAI, "
        "2023), which injects random strings into the system context and monitors outputs for "
        "their presence. However, canary tokens only detect successful attacks that result in "
        "context leakage \u2014 they do not prevent the injection itself. Sophisticated attackers "
        "can instruct the model to strip unusual strings from output, and the approach is "
        "limited to detecting prompt leaking rather than goal hijacking or tool manipulation "
        "(Abdallah et al., 2025)."
    ))

    # LLM-as-Judge - with citations
    doc.add_heading("LLM-as-Judge Detection", level=3)
    _add_body_paragraph(doc, DETECTION_TECHNIQUES[4]["description"])
    _add_body_paragraph(doc, (
        "This model-based approach is exemplified by the Rebuff LLM-based detection layer "
        "(ProtectAI, 2023) and dual-LLM architectures where a smaller model screens inputs "
        "before the primary model processes them (Abdallah et al., 2025). However, this adds "
        "latency and cost from a second model inference per request. The judge model is itself "
        "susceptible to prompt injection, and effectiveness degrades when the judge model is "
        "weaker than the primary model (Yi et al., 2024)."
    ))

    # Comparison Table
    _add_apa_table(
        doc,
        number=1,
        title="Comparison of Prompt Injection Detection Techniques",
        headers=["Technique", "Type", "Speed", "Known Patterns", "Novel Attacks", "Key Limitation"],
        rows=[
            ["Heuristic/Rule-Based", "Static", "High", "Moderate", "Low", "Trivially bypassed via paraphrasing"],
            ["ML Classification", "Learned", "Moderate", "High", "Low\u2013Moderate", "Training data dependent"],
            ["Perplexity Analysis", "Statistical", "Moderate", "Low", "High (machine)", "Ineffective on human-crafted"],
            ["Canary Tokens", "Proactive", "High", "N/A (post-hoc)", "N/A", "Only detects context leakage"],
            ["LLM-as-Judge", "Model-based", "Low", "High", "Moderate\u2013High", "Costly; itself vulnerable"],
        ]
    )

    # Embed Figure 2 after detection
    _embed_figure(doc, 1, 2)

    _add_body_paragraph(doc, (
        "Figure 2 illustrates the detection and defense pipeline, showing how multiple "
        "detection layers can be composed in sequence to improve overall detection rates. "
        "The multi-layered approach reflects the defense-in-depth principle: no single "
        "detection method provides sufficient coverage, but their combination significantly "
        "reduces the probability of a successful attack evading all layers."
    ))


def _add_prevention_subsection(doc: Document):
    """Prevention Strategies subsection with comparison table."""
    print("  Writing subsection: Prevention Strategies...")
    doc.add_heading("Prevention Strategies", level=2)

    _add_body_paragraph(doc, (
        "Prevention strategies for prompt injection operate at three levels: model-level "
        "defenses that modify how the LLM processes instructions, application-level controls "
        "that filter inputs and outputs, and system-level architectural patterns that limit "
        "the blast radius of successful attacks. This section evaluates six major prevention "
        "approaches and their practical tradeoffs."
    ))

    for strategy in PREVENTION_STRATEGIES:
        doc.add_heading(strategy["strategy"], level=3)
        _add_body_paragraph(doc, strategy["description"])
        _add_body_paragraph(doc, (
            f"Operating at the {strategy['implementation_level']} level, this strategy is "
            f"implemented through approaches such as {strategy['tools_examples'][0]}. "
            f"{strategy['effectiveness']}"
        ))

    # Comparison Table
    _add_apa_table(
        doc,
        number=2,
        title="Comparison of Prompt Injection Prevention Strategies",
        headers=["Strategy", "Level", "Prevents Injection", "Limits Damage", "Complexity", "Coverage"],
        rows=[
            ["Instruction Hierarchy", "Model", "High (direct)", "N/A", "High", "Partial (indirect bypasses)"],
            ["Input Sanitization", "Application", "Moderate", "N/A", "Low", "Known patterns only"],
            ["Output Validation", "Application", "N/A (post-hoc)", "Moderate", "Low", "Observable violations"],
            ["Sandboxing/Isolation", "System", "None", "High", "Moderate", "Broad blast-radius control"],
            ["Guardrail Frameworks", "Application", "Moderate", "Moderate", "Moderate", "Configurable per app"],
            ["Architectural Patterns", "System", "High (theoretical)", "High", "High", "Broad but reduces utility"],
        ]
    )

    # Embed Figure 3 after prevention
    _embed_figure(doc, 2, 3)

    _add_body_paragraph(doc, (
        "Figure 3 presents the system architecture for a layered prompt injection defense, "
        "illustrating how model-level, application-level, and system-level controls integrate "
        "to form a comprehensive defense posture. The architecture emphasizes that effective "
        "defense requires coordination across all three levels, as no single layer provides "
        "complete protection against the full spectrum of attack vectors."
    ))


def _embed_figure(doc: Document, index: int, figure_number: int):
    """Embed a figure with APA 7th formatted caption.

    APA 7th: Bold figure number on its own line, italic title on next line.
    """
    filename, caption_text = FIGURE_FILES[index]
    figure_path = FIGURES_DIR / filename

    if not figure_path.exists():
        raise FileNotFoundError(
            f"Figure not found: {figure_path}. "
            f"Ensure results/figures/{filename} exists before generating the paper."
        )

    # Add the image centered
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(figure_path), width=Inches(5.5))

    # Figure number (bold, flush left per APA 7th)
    num_para = doc.add_paragraph()
    num_para.paragraph_format.space_before = Pt(6)
    num_para.paragraph_format.space_after = Pt(0)
    num_para.paragraph_format.first_line_indent = Inches(0)
    run = num_para.add_run(f"Figure {figure_number}")
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    # Caption title (italic)
    caption_para = doc.add_paragraph()
    caption_para.paragraph_format.space_after = Pt(12)
    caption_para.paragraph_format.first_line_indent = Inches(0)
    run = caption_para.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"


def _add_methodology(doc: Document):
    """Add the Methodology section."""
    print("Writing section: Methodology...")
    doc.add_heading("Methodology", level=1)

    paras = [
        (
            "This research employs a structured literature survey methodology to examine "
            "the current state of prompt injection attacks and defenses. The survey follows "
            "a systematic approach to source selection, categorization, and synthesis designed "
            "to provide comprehensive coverage while maintaining analytical rigor."
        ),
        (
            "Source selection targeted four categories of publications: peer-reviewed academic "
            "papers from arXiv and journals, industry reports from major cybersecurity firms "
            "and AI companies, government standards and frameworks, and documentation from "
            "open-source security tools. A total of 18 primary sources published between 2023 "
            "and 2025 were selected based on relevance, credibility, and recency. Academic "
            "sources were prioritized for theoretical foundations and empirical findings, while "
            "industry and government sources provided practical context and standardized "
            "frameworks."
        ),
        (
            "The analysis framework organizes findings across four pillars: attack taxonomy "
            "and classification, risk analysis and impact assessment, detection techniques, "
            "and prevention strategies. Each pillar draws from multiple sources to enable "
            "cross-validation and identify areas of consensus and disagreement. The taxonomy "
            "was constructed through iterative synthesis, beginning with established categories "
            "from Greshake et al. (2023) and extending to incorporate emerging threat vectors "
            "documented in 2024-2025 publications."
        ),
        (
            "Limitations of this methodology include publication bias toward well-documented "
            "attack vectors, the rapid pace of development in the field which may render some "
            "findings outdated, and the difficulty of verifying effectiveness claims for "
            "defensive tools in the absence of standardized benchmarks. The survey scope is "
            "limited to prompt injection specifically and does not cover other LLM security "
            "concerns such as training data poisoning or model extraction attacks."
        ),
    ]
    for text in paras:
        _add_body_paragraph(doc, text)


def _add_discussion(doc: Document):
    """Add the Discussion section with flowing analysis."""
    print("Writing section: Discussion...")
    doc.add_heading("Discussion", level=1)

    paras = [
        (
            "The synthesis of findings reveals that the evolution of prompt injection directly "
            "mirrors the expansion of LLM capabilities. As language models gained tool-use, web "
            "browsing, and multi-agent orchestration capabilities between 2022 and 2025, the "
            "attack taxonomy expanded correspondingly \u2014 from simple direct injection techniques "
            "to the autonomous propagating threats documented by Ahmed et al. (2025). This "
            "co-evolution is not coincidental; each new capability creates a new attack surface "
            "that adversaries rapidly exploit. The trajectory suggests that future LLM capabilities "
            "will inevitably introduce novel injection vectors, making static defense approaches "
            "fundamentally insufficient."
        ),
        (
            "The relationship between detection limitations and risk severity warrants particular "
            "attention. Current detection architectures \u2014 even multi-layered approaches combining "
            "heuristic, ML-based, and LLM-as-judge methods \u2014 face a structural disadvantage: "
            "novel injection techniques can be crafted more quickly than detection systems can "
            "adapt (Abdallah et al., 2025). This asymmetry directly enables the most severe risk "
            "categories. Data exfiltration through markdown image encoding (Greshake et al., 2023) "
            "and unauthorized tool invocation through indirect injection (OWASP Foundation, 2025) "
            "succeed precisely because they exploit gaps between detection layers. The parallel to "
            "the antivirus industry's evolution from signature-based to behavioral detection is "
            "instructive \u2014 the prompt injection field appears to be in its signature-based era, "
            "where reactive detection provides inadequate security against motivated adversaries."
        ),
        (
            "Prevention strategy analysis reveals a fundamental tension between security and "
            "utility that divides responsibility across stakeholders. The instruction hierarchy "
            "approach (Wallace et al., 2024) represents the most theoretically sound defense, "
            "as it addresses the root cause \u2014 the inability to distinguish instruction privilege "
            "levels \u2014 at the model architecture level. However, this approach is exclusively "
            "available to model providers, leaving application developers reliant on input "
            "sanitization, guardrail frameworks, and output validation as their primary defenses. "
            "System architects can complement these with sandboxing and least-privilege patterns "
            "that limit blast radius, but at the cost of reduced LLM autonomy. This three-tier "
            "division \u2014 model providers handling instruction hierarchy, developers deploying "
            "guardrails, and architects enforcing isolation \u2014 represents the emerging consensus "
            "on defense allocation."
        ),
        (
            "The supply chain and propagation risks compound these defensive challenges. The "
            "Morris II AI worm concept demonstrates that prompt injection is not merely a point "
            "vulnerability but a potential vector for systemic cascading failure across "
            "interconnected AI systems. When the output of a compromised model becomes input to "
            "another \u2014 through shared knowledge bases, AI-generated code repositories, or "
            "multi-agent workflows \u2014 a single injection can propagate beyond organizational "
            "boundaries (Ahmed et al., 2025). This systemic risk dimension distinguishes prompt "
            "injection from traditional web application vulnerabilities and demands architectural "
            "defenses, particularly the dual-LLM pattern and human-in-the-loop confirmation, "
            "despite their impact on system autonomy."
        ),
        (
            "The implications for practitioners are clear: organizations deploying LLM-integrated "
            "systems must adopt layered defense strategies that combine multiple prevention "
            "mechanisms rather than relying on any single mitigation. The principle of least "
            "privilege should govern all tool access granted to LLM agents. Human-in-the-loop "
            "confirmation should be required for irreversible or high-impact operations. Security "
            "teams should treat prompt injection with the same rigor applied to traditional web "
            "application vulnerabilities \u2014 including penetration testing against injection "
            "vectors, continuous monitoring for anomalous model behavior, and incident response "
            "planning that accounts for the unique characteristics of AI-mediated attacks."
        ),
    ]
    for text in paras:
        _add_body_paragraph(doc, text)


def _add_conclusion(doc: Document):
    """Add the Conclusion section with specific future direction citations."""
    print("Writing section: Conclusion...")
    doc.add_heading("Conclusion", level=1)

    paras = [
        (
            "This survey has examined prompt injection attacks across four dimensions: "
            "taxonomy, risks, detection, and prevention. The analysis reveals that prompt "
            "injection represents a fundamental and largely unsolved challenge in LLM security, "
            "rooted in the architectural inability of current language models to reliably "
            "distinguish between instructions and data within their context window."
        ),
        (
            "The attack taxonomy has grown from simple direct injection techniques to encompass "
            "indirect injection through external data sources, multimodal attacks exploiting "
            "vision and audio channels, tool and agent-based exploitation of agentic AI systems, "
            "hybrid chained attacks that combine multiple vectors, and autonomous propagating "
            "threats that can spread without continued attacker involvement. Each new category "
            "reflects the expanding capabilities and deployment contexts of modern LLMs."
        ),
        (
            "The current state of detection and prevention offers practical defenses but no "
            "complete solutions. Defense-in-depth strategies combining instruction hierarchy, "
            "input and output filtering, guardrail frameworks, and architectural sandboxing "
            "provide meaningful risk reduction. However, the fundamental vulnerability persists "
            "as an inherent property of the transformer architecture's treatment of all context "
            "window content as equally authoritative."
        ),
        (
            "Future research directions include the development of formal verification methods "
            "for LLM instruction compliance, as preliminary work by Liu et al. (2024) on "
            "benchmarking suggests that standardized evaluation frameworks are both feasible "
            "and urgently needed. The dual-LLM architectural pattern warrants further "
            "investigation as a practical defense for production deployments, particularly in "
            "establishing the optimal relationship between quarantined and privileged models. "
            "Additionally, the extension of instruction hierarchy training (Wallace et al., "
            "2024) to handle indirect injection through trusted data channels remains a critical "
            "open problem. Until such advances materialize, the cybersecurity community must "
            "treat prompt injection as a persistent threat requiring continuous monitoring, "
            "testing, and adaptive defense strategies."
        ),
    ]
    for text in paras:
        _add_body_paragraph(doc, text)


def _add_references(doc: Document):
    """Add the References section with APA hanging indent."""
    print("Writing section: References...")
    doc.add_page_break()
    doc.add_heading("References", level=1)

    for ref in REFERENCES:
        para = doc.add_paragraph(ref)
        # Hanging indent: 0.5 inch left indent, -0.5 inch first line (net: first line at 0)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.paragraph_format.space_after = Pt(0)


def generate_paper():
    """Generate the complete research paper and write to output/research_paper.docx."""
    print("Generating research paper...")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / "research_paper.docx"

    doc = Document()
    _set_apa_defaults(doc)
    _add_page_numbers(doc)
    _add_title_page(doc)
    _add_abstract(doc)
    _add_introduction(doc)
    _add_literature_review(doc)
    _add_methodology(doc)
    _add_discussion(doc)
    _add_conclusion(doc)
    _add_references(doc)

    doc.save(str(output_path))
    print(f"Paper saved to {output_path}")
    print(f"File size: {output_path.stat().st_size:,} bytes")
    return output_path


if __name__ == "__main__":
    generate_paper()
